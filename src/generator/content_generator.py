"""
Content generator with RAG pipeline and brand voice integration.
Production-ready implementation for Erguvan AI Content Generator.
"""

import json
import logging
import asyncio
from typing import Dict, List, Any, Optional, Tuple
from dataclasses import dataclass
from datetime import datetime
from pathlib import Path
import re

from openai import OpenAI
from anthropic import Anthropic
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

from src.config import config_manager
from src.prompts import prompt_manager
from src.vector_index.vector_store import VectorStoreManager, SearchResult
from src.analyzer.style_analyzer import StyleProfileManager


@dataclass
class ContentRequest:
    """Request for content generation."""
    topic: str
    audience: str
    desired_length: int
    language: str = "en"
    style_override: Optional[str] = None
    additional_context: Optional[str] = None
    
    def __post_init__(self):
        """Validate request parameters."""
        if self.desired_length < 100:
            raise ValueError("Desired length must be at least 100 words")
        if self.desired_length > 10000:
            raise ValueError("Desired length cannot exceed 10,000 words")
        if self.language not in ["en", "tr"]:
            raise ValueError("Language must be 'en' or 'tr'")


@dataclass
class ContentSection:
    """Individual content section."""
    heading: str
    body: str
    word_count: int
    
    def __post_init__(self):
        """Calculate word count if not provided."""
        if not self.word_count:
            self.word_count = len(self.body.split())


@dataclass
class Citation:
    """Citation information."""
    source: str
    url: Optional[str] = None
    relevance: str = ""
    document_id: Optional[str] = None


@dataclass
class GeneratedContent:
    """Generated content with metadata."""
    title: str
    author: str
    created_utc: datetime
    audience: str
    word_count: int
    language: str
    sections: List[ContentSection]
    citations: List[Citation]
    key_takeaways: List[str]
    
    # Generation metadata
    topic: str
    style_patterns: List[str]
    context_chunks_used: int
    generation_time_seconds: float
    model_used: str
    
    def to_dict(self) -> Dict[str, Any]:
        """Convert to dictionary for JSON serialization."""
        return {
            "title": self.title,
            "author": self.author,
            "created_utc": self.created_utc.isoformat(),
            "audience": self.audience,
            "word_count": self.word_count,
            "language": self.language,
            "sections": [
                {
                    "heading": section.heading,
                    "body": section.body,
                    "word_count": section.word_count
                }
                for section in self.sections
            ],
            "citations": [
                {
                    "source": citation.source,
                    "url": citation.url,
                    "relevance": citation.relevance,
                    "document_id": citation.document_id
                }
                for citation in self.citations
            ],
            "key_takeaways": self.key_takeaways,
            "topic": self.topic,
            "style_patterns": self.style_patterns,
            "context_chunks_used": self.context_chunks_used,
            "generation_time_seconds": self.generation_time_seconds,
            "model_used": self.model_used
        }


class LLMContentGenerator:
    """LLM-powered content generation."""
    
    def __init__(self):
        self.logger = logging.getLogger(__name__)
        self.config = config_manager.get_model_config()
        
        # Initialize LLM client
        if self.config.provider == "openai":
            self.client = OpenAI()
        elif self.config.provider == "anthropic":
            self.client = Anthropic()
        else:
            raise ValueError(f"Unsupported provider: {self.config.provider}")
    
    async def generate_content(self, prompt: str, max_tokens: int = 4000) -> str:
        """Generate content using LLM with retry logic."""
        max_retries = 3
        
        for attempt in range(max_retries):
            try:
                if self.config.provider == "openai":
                    response = self.client.chat.completions.create(
                        model=self.config.generation_model,
                        messages=[{"role": "user", "content": prompt}],
                        temperature=self.config.temperature,
                        max_tokens=max_tokens,
                        top_p=0.9
                    )
                    return response.choices[0].message.content
                
                elif self.config.provider == "anthropic":
                    response = self.client.messages.create(
                        model=self.config.generation_model,
                        messages=[{"role": "user", "content": prompt}],
                        temperature=self.config.temperature,
                        max_tokens=max_tokens
                    )
                    return response.content[0].text
                
            except Exception as e:
                self.logger.warning(f"LLM generation attempt {attempt + 1} failed: {e}")
                if attempt == max_retries - 1:
                    raise
                await asyncio.sleep(2 ** attempt)  # Exponential backoff
    
    def estimate_cost(self, prompt_length: int, response_length: int) -> float:
        """Estimate generation cost based on token usage."""
        # Rough estimation: 1 token ≈ 0.75 words
        prompt_tokens = prompt_length * 0.75
        response_tokens = response_length * 0.75
        
        if self.config.provider == "openai":
            if "gpt-4" in self.config.generation_model:
                # GPT-4 pricing (approximate)
                input_cost = prompt_tokens * 0.03 / 1000
                output_cost = response_tokens * 0.06 / 1000
                return input_cost + output_cost
            else:
                # GPT-3.5 pricing (approximate)
                return (prompt_tokens + response_tokens) * 0.002 / 1000
        
        # Default estimation
        return 0.05


class RAGPipeline:
    """Retrieval-Augmented Generation pipeline."""
    
    def __init__(self, vector_store_manager: VectorStoreManager):
        self.vector_store = vector_store_manager
        self.logger = logging.getLogger(__name__)
    
    async def retrieve_context(self, request: ContentRequest) -> List[SearchResult]:
        """Retrieve relevant context for content generation."""
        self.logger.info(f"Retrieving context for topic: {request.topic}")
        
        # Search for relevant chunks
        context_chunks = await self.vector_store.search_for_generation(
            topic=request.topic,
            audience=request.audience,
            language=request.language,
            max_chunks=8  # Get more chunks for better context
        )
        
        # Filter and rank chunks
        filtered_chunks = self._filter_and_rank_chunks(context_chunks, request)
        
        self.logger.info(f"Retrieved {len(filtered_chunks)} context chunks")
        
        return filtered_chunks
    
    def _filter_and_rank_chunks(self, chunks: List[SearchResult], 
                               request: ContentRequest) -> List[SearchResult]:
        """Filter and rank chunks based on relevance and quality."""
        if not chunks:
            return []
        
        # Filter by minimum relevance score
        min_relevance = 0.3
        filtered = [chunk for chunk in chunks if chunk.relevance_score >= min_relevance]
        
        # Re-rank based on multiple factors
        scored_chunks = []
        for chunk in filtered:
            score = self._calculate_chunk_score(chunk, request)
            scored_chunks.append((chunk, score))
        
        # Sort by score and return top 5
        scored_chunks.sort(key=lambda x: x[1], reverse=True)
        
        return [chunk for chunk, score in scored_chunks[:5]]
    
    def _calculate_chunk_score(self, chunk: SearchResult, 
                              request: ContentRequest) -> float:
        """Calculate composite score for chunk ranking."""
        score = chunk.relevance_score
        
        # Boost score for audience match
        if request.audience.lower() in chunk.content.lower():
            score *= 1.2
        
        # Boost score for topic keywords
        topic_keywords = request.topic.lower().split()
        content_lower = chunk.content.lower()
        keyword_matches = sum(1 for keyword in topic_keywords if keyword in content_lower)
        score *= (1 + keyword_matches * 0.1)
        
        # Boost score for appropriate chunk length
        word_count = len(chunk.content.split())
        if 100 <= word_count <= 500:  # Ideal chunk size
            score *= 1.1
        
        # Boost score for chunks with headings
        if chunk.heading:
            score *= 1.05
        
        return score


class StylePatternExtractor:
    """Extracts style patterns from similar documents."""
    
    def __init__(self, style_manager: StyleProfileManager):
        self.style_manager = style_manager
        self.logger = logging.getLogger(__name__)
    
    async def extract_patterns(self, context_chunks: List[SearchResult], 
                             request: ContentRequest) -> List[str]:
        """Extract style patterns from context chunks."""
        patterns = []
        
        # Get document IDs from context chunks
        doc_ids = list(set(chunk.document_id for chunk in context_chunks))
        
        # Get aggregated style patterns
        if doc_ids:
            aggregate_style = self.style_manager.get_aggregate_style(doc_ids)
            
            if aggregate_style:
                # Convert style data to pattern descriptions
                patterns.extend(self._style_to_patterns(aggregate_style))
        
        # Add audience-specific patterns
        patterns.extend(self._get_audience_patterns(request.audience))
        
        # Add language-specific patterns
        patterns.extend(self._get_language_patterns(request.language))
        
        self.logger.info(f"Extracted {len(patterns)} style patterns")
        
        return patterns
    
    def _style_to_patterns(self, style_data: Dict[str, Any]) -> List[str]:
        """Convert style data to pattern descriptions."""
        patterns = []
        
        if 'avg_sentence_length' in style_data:
            length = style_data['avg_sentence_length']
            if length < 15:
                patterns.append("Use short, concise sentences")
            elif length > 25:
                patterns.append("Use longer, more detailed sentences")
            else:
                patterns.append("Use medium-length sentences for clarity")
        
        if 'avg_formality_score' in style_data:
            formality = style_data['avg_formality_score']
            if formality > 7:
                patterns.append("Maintain formal, professional tone")
            elif formality < 4:
                patterns.append("Use conversational, accessible language")
            else:
                patterns.append("Balance formal and conversational elements")
        
        if 'common_phrases' in style_data:
            phrases = style_data['common_phrases'][:3]
            if phrases:
                patterns.append(f"Include relevant terms: {', '.join(phrases)}")
        
        if 'sentence_starters' in style_data:
            starters = style_data['sentence_starters'][:3]
            if starters:
                patterns.append(f"Use varied sentence openings like: {', '.join(starters)}")
        
        return patterns
    
    def _get_audience_patterns(self, audience: str) -> List[str]:
        """Get style patterns based on target audience."""
        audience_lower = audience.lower()
        
        if "cfo" in audience_lower or "executive" in audience_lower:
            return [
                "Focus on strategic implications and ROI",
                "Use executive summary format",
                "Emphasize quantified benefits and risks",
                "Include actionable recommendations"
            ]
        elif "sustainability" in audience_lower or "esg" in audience_lower:
            return [
                "Include technical details and methodologies",
                "Reference relevant frameworks and standards",
                "Provide comprehensive coverage of topic",
                "Use industry-specific terminology"
            ]
        elif "procurement" in audience_lower or "manager" in audience_lower:
            return [
                "Focus on operational implications",
                "Provide practical implementation guidance",
                "Include timelines and process steps",
                "Emphasize supply chain impacts"
            ]
        else:
            return [
                "Use clear, accessible language",
                "Provide balanced coverage of topic",
                "Include relevant examples and context"
            ]
    
    def _get_language_patterns(self, language: str) -> List[str]:
        """Get style patterns based on language."""
        if language == "tr":
            return [
                "Use Turkish sustainability terminology",
                "Reference Turkish and EU regulations",
                "Adapt sentence structure for Turkish grammar",
                "Include cultural context for Turkish audience"
            ]
        else:
            return [
                "Use international sustainability standards",
                "Reference global regulatory frameworks",
                "Maintain clear, direct English style",
                "Include international best practices"
            ]


class ContentStructurer:
    """Structures generated content into proper sections."""
    
    def __init__(self):
        self.logger = logging.getLogger(__name__)
    
    def structure_content(self, raw_content: str, request: ContentRequest) -> List[ContentSection]:
        """Structure raw content into logical sections."""
        try:
            # Try to parse as JSON first
            content_data = json.loads(raw_content)
            if 'sections' in content_data:
                sections = []
                for section_data in content_data['sections']:
                    section = ContentSection(
                        heading=section_data.get('heading', ''),
                        body=section_data.get('body', ''),
                        word_count=0  # Will be calculated in __post_init__
                    )
                    sections.append(section)
                return sections
        except json.JSONDecodeError:
            pass
        
        # Fallback: structure plain text
        return self._structure_plain_text(raw_content, request)
    
    def _structure_plain_text(self, content: str, request: ContentRequest) -> List[ContentSection]:
        """Structure plain text content into sections."""
        sections = []
        
        # Split by common heading patterns
        lines = content.split('\n')
        current_heading = None
        current_body = []
        
        for line in lines:
            line = line.strip()
            if not line:
                continue
            
            # Check if line is a heading
            if self._is_heading(line):
                # Save previous section
                if current_heading and current_body:
                    section = ContentSection(
                        heading=current_heading,
                        body='\n'.join(current_body).strip(),
                        word_count=0
                    )
                    sections.append(section)
                
                # Start new section
                current_heading = line
                current_body = []
            else:
                current_body.append(line)
        
        # Add final section
        if current_heading and current_body:
            section = ContentSection(
                heading=current_heading,
                body='\n'.join(current_body).strip(),
                word_count=0
            )
            sections.append(section)
        
        # If no sections found, create single section
        if not sections:
            sections = [ContentSection(
                heading="Content",
                body=content.strip(),
                word_count=0
            )]
        
        return sections
    
    def _is_heading(self, line: str) -> bool:
        """Check if a line is likely a heading."""
        # Common heading patterns
        patterns = [
            r'^#{1,6}\s+',  # Markdown headings
            r'^[A-Z][A-Za-z\s]+:$',  # Colon-ending headings
            r'^\d+\.\s+[A-Z]',  # Numbered headings
            r'^[A-Z][A-Z\s]+$',  # ALL CAPS (short lines only)
        ]
        
        for pattern in patterns:
            if re.match(pattern, line):
                return True
        
        # Check for ALL CAPS headings (must be reasonably short)
        if line.isupper() and len(line) < 80 and len(line.split()) <= 8:
            return True
        
        return False


class ContentGenerator:
    """Main content generator orchestrating all components."""
    
    def __init__(self, vector_store_manager: VectorStoreManager, 
                 style_manager: StyleProfileManager):
        self.vector_store = vector_store_manager
        self.style_manager = style_manager
        self.llm_generator = LLMContentGenerator()
        self.rag_pipeline = RAGPipeline(vector_store_manager)
        self.style_extractor = StylePatternExtractor(style_manager)
        self.structurer = ContentStructurer()
        self.logger = logging.getLogger(__name__)
    
    async def generate_content(self, request: ContentRequest) -> GeneratedContent:
        """Generate content based on request."""
        start_time = datetime.now()
        
        self.logger.info(f"Generating content for topic: {request.topic}")
        
        try:
            # Step 1: Retrieve context
            context_chunks = await self.rag_pipeline.retrieve_context(request)
            
            # Step 2: Extract style patterns
            style_patterns = await self.style_extractor.extract_patterns(context_chunks, request)
            
            # Step 3: Generate content prompt
            prompt = prompt_manager.render_content_generation_prompt(
                topic=request.topic,
                audience=request.audience,
                desired_length=request.desired_length,
                context_chunks=context_chunks,
                style_patterns=style_patterns,
                language=request.language,
                style_override=request.style_override
            )
            
            # Step 4: Generate content
            raw_content = await self.llm_generator.generate_content(prompt)
            
            # Step 5: Structure content
            sections = self.structurer.structure_content(raw_content, request)
            
            # Step 6: Extract additional information
            title, citations, key_takeaways = self._extract_metadata(raw_content, context_chunks)
            
            # Step 7: Calculate metrics
            total_word_count = sum(section.word_count for section in sections)
            generation_time = (datetime.now() - start_time).total_seconds()
            
            # Step 8: Create content object
            content = GeneratedContent(
                title=title,
                author="Erguvan Advisory AI",
                created_utc=datetime.utcnow(),
                audience=request.audience,
                word_count=total_word_count,
                language=request.language,
                sections=sections,
                citations=citations,
                key_takeaways=key_takeaways,
                topic=request.topic,
                style_patterns=style_patterns,
                context_chunks_used=len(context_chunks),
                generation_time_seconds=generation_time,
                model_used=self.llm_generator.config.generation_model
            )
            
            self.logger.info(f"Content generation completed in {generation_time:.2f}s "
                           f"({total_word_count} words)")
            
            return content
            
        except Exception as e:
            self.logger.error(f"Content generation failed: {e}")
            raise
    
    def _extract_metadata(self, raw_content: str, 
                         context_chunks: List[SearchResult]) -> Tuple[str, List[Citation], List[str]]:
        """Extract title, citations, and key takeaways from content."""
        # Try to parse as JSON first
        try:
            content_data = json.loads(raw_content)
            
            title = content_data.get('title', 'Generated Content')
            
            # Extract citations
            citations = []
            for citation_data in content_data.get('citations', []):
                citation = Citation(
                    source=citation_data.get('source', ''),
                    url=citation_data.get('url'),
                    relevance=citation_data.get('relevance', '')
                )
                citations.append(citation)
            
            # Extract key takeaways
            key_takeaways = content_data.get('key_takeaways', [])
            
            return title, citations, key_takeaways
            
        except json.JSONDecodeError:
            pass
        
        # Fallback: extract from plain text
        title = self._extract_title_from_text(raw_content)
        citations = self._create_citations_from_chunks(context_chunks)
        key_takeaways = self._extract_key_points_from_text(raw_content)
        
        return title, citations, key_takeaways
    
    def _extract_title_from_text(self, content: str) -> str:
        """Extract title from plain text content."""
        lines = content.split('\n')
        for line in lines:
            line = line.strip()
            if line and not line.startswith('#'):
                # Use first substantial line as title
                if len(line) > 10 and len(line) < 100:
                    return line
        
        return "Generated Content"
    
    def _create_citations_from_chunks(self, chunks: List[SearchResult]) -> List[Citation]:
        """Create citations from context chunks."""
        citations = []
        
        for chunk in chunks:
            citation = Citation(
                source=f"Document: {chunk.metadata.get('file_name', 'Unknown')}",
                relevance=f"Relevance score: {chunk.relevance_score:.2f}",
                document_id=chunk.document_id
            )
            citations.append(citation)
        
        return citations
    
    def _extract_key_points_from_text(self, content: str) -> List[str]:
        """Extract key takeaways from plain text."""
        # Simple extraction of bullet points or numbered lists
        lines = content.split('\n')
        key_points = []
        
        for line in lines:
            line = line.strip()
            if re.match(r'^[•\-\*]\s+', line) or re.match(r'^\d+\.\s+', line):
                point = re.sub(r'^[•\-\*\d\.]\s+', '', line)
                if len(point) > 20:  # Substantial points only
                    key_points.append(point)
        
        # Limit to top 5 key points
        return key_points[:5]
    
    async def save_to_docx(self, content: GeneratedContent, output_path: Path) -> None:
        """Save generated content to DOCX file."""
        try:
            # Create Word document
            doc = Document()
            
            # Add title
            title_paragraph = doc.add_heading(content.title, 0)
            title_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Add metadata
            doc.add_paragraph(f"Author: {content.author}")
            doc.add_paragraph(f"Created: {content.created_utc.strftime('%Y-%m-%d %H:%M:%S')} UTC")
            doc.add_paragraph(f"Audience: {content.audience}")
            doc.add_paragraph(f"Word Count: {content.word_count:,}")
            doc.add_paragraph(f"Language: {content.language.upper()}")
            
            # Add separator
            doc.add_page_break()
            
            # Add sections
            for section in content.sections:
                if section.heading:
                    doc.add_heading(section.heading, 1)
                
                # Add body paragraphs
                for paragraph in section.body.split('\n\n'):
                    if paragraph.strip():
                        doc.add_paragraph(paragraph.strip())
            
            # Add key takeaways
            if content.key_takeaways:
                doc.add_heading("Key Takeaways", 1)
                for takeaway in content.key_takeaways:
                    doc.add_paragraph(takeaway, style='List Bullet')
            
            # Add citations
            if content.citations:
                doc.add_heading("References", 1)
                for i, citation in enumerate(content.citations, 1):
                    citation_text = f"{i}. {citation.source}"
                    if citation.url:
                        citation_text += f" ({citation.url})"
                    doc.add_paragraph(citation_text)
            
            # Save document
            doc.save(output_path)
            
            self.logger.info(f"Content saved to DOCX: {output_path}")
            
        except Exception as e:
            self.logger.error(f"Failed to save content to DOCX: {e}")
            raise