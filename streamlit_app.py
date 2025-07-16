"""
Streamlit UI for Erguvan AI Content Generator
Simple web interface for editor review workflow.
"""

import streamlit as st
import asyncio
import json
import os
import logging
from pathlib import Path
from datetime import datetime
import pandas as pd
from dotenv import load_dotenv
from concurrent.futures import ThreadPoolExecutor
import threading

# Load environment variables
load_dotenv()

# Setup logging
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s',
    handlers=[
        logging.StreamHandler(),
        logging.FileHandler('streamlit_app.log', mode='a')
    ]
)
logger = logging.getLogger(__name__)

# Add src to path with better error handling
import sys
src_path = Path(__file__).parent / "src"
if src_path.exists():
    sys.path.insert(0, str(src_path))
    logger.info(f"Added {src_path} to Python path")
else:
    logger.error(f"Source directory not found: {src_path}")
    st.error(f"Source directory not found: {src_path}")

# Import with error handling
try:
    from src.config import config_manager
    # Try to import MarkItDown first, fall back to original DocumentLoader
    try:
        from src.loader.markitdown_loader import MarkItDownLoader, DocumentBatch
        MARKITDOWN_AVAILABLE = True
        logger.info("MarkItDown loader is available")
    except ImportError as e:
        logger.warning(f"MarkItDown not available, falling back to DocumentLoader: {e}")
        from src.loader.document_loader import DocumentLoader, DocumentBatch
        MARKITDOWN_AVAILABLE = False
    from src.analyzer.style_analyzer import StyleProfileManager
    from src.vector_index.vector_store import VectorStoreManager
    from src.generator.content_generator import ContentGenerator, ContentRequest
    from src.evaluator.quality_evaluator import QualityEvaluator
    logger.info("Successfully imported all source modules")
except ImportError as e:
    logger.error(f"Failed to import source modules: {e}")
    st.error(f"Failed to import required modules: {e}")
    st.stop()


# Configure Streamlit
st.set_page_config(
    page_title="Erguvan AI Content Generator",
    page_icon="🌱",
    layout="wide",
    initial_sidebar_state="expanded"
)

# Environment validation
def validate_environment():
    """Validate environment and configuration files."""
    issues = []
    
    # Check environment variables
    required_env_vars = ['OPENAI_API_KEY']
    for var in required_env_vars:
        if not os.getenv(var):
            issues.append(f"Missing environment variable: {var}")
    
    # Check config files
    config_dir = Path(__file__).parent / "config"
    required_configs = ['models.yaml', 'brand_guide.yaml']
    for config_file in required_configs:
        config_path = config_dir / config_file
        if not config_path.exists():
            issues.append(f"Missing config file: {config_path}")
    
    # Check data directory
    data_dir = Path(__file__).parent / "data"
    if not data_dir.exists():
        logger.warning(f"Data directory does not exist: {data_dir}")
        data_dir.mkdir(parents=True, exist_ok=True)
        logger.info(f"Created data directory: {data_dir}")
    
    return issues

# Initialize session state
if 'vector_store_manager' not in st.session_state:
    st.session_state.vector_store_manager = None
if 'style_manager' not in st.session_state:
    st.session_state.style_manager = None
if 'generator' not in st.session_state:
    st.session_state.generator = None
if 'evaluator' not in st.session_state:
    st.session_state.evaluator = None
if 'initialization_errors' not in st.session_state:
    st.session_state.initialization_errors = []
if 'async_executor' not in st.session_state:
    st.session_state.async_executor = ThreadPoolExecutor(max_workers=2)


# Async helper functions
def run_async_in_thread(coro):
    """Run async function in thread pool to avoid event loop conflicts."""
    def run_in_thread():
        try:
            loop = asyncio.new_event_loop()
            asyncio.set_event_loop(loop)
            return loop.run_until_complete(coro)
        finally:
            loop.close()
    
    return st.session_state.async_executor.submit(run_in_thread).result()

@st.cache_resource
def initialize_components():
    """Initialize system components with improved error handling."""
    errors = []
    
    try:
        logger.info("Initializing vector store manager...")
        vector_store_manager = VectorStoreManager()
        logger.info("Vector store manager initialized successfully")
    except Exception as e:
        error_msg = f"Failed to initialize vector store manager: {e}"
        logger.error(error_msg)
        errors.append(error_msg)
        vector_store_manager = None
    
    try:
        logger.info("Initializing style manager...")
        style_manager = StyleProfileManager()
        
        # Load existing style profiles if available
        style_profiles_path = Path("data/style_profiles.json")
        if style_profiles_path.exists():
            logger.info(f"Loading style profiles from {style_profiles_path}")
            style_manager.load_profiles(str(style_profiles_path))
        else:
            logger.warning(f"Style profiles file not found: {style_profiles_path}")
        
        logger.info("Style manager initialized successfully")
    except Exception as e:
        error_msg = f"Failed to initialize style manager: {e}"
        logger.error(error_msg)
        errors.append(error_msg)
        style_manager = None
    
    try:
        if vector_store_manager and style_manager:
            logger.info("Initializing content generator...")
            generator = ContentGenerator(vector_store_manager, style_manager)
            logger.info("Content generator initialized successfully")
        else:
            generator = None
            errors.append("Cannot initialize content generator: missing dependencies")
    except Exception as e:
        error_msg = f"Failed to initialize content generator: {e}"
        logger.error(error_msg)
        errors.append(error_msg)
        generator = None
    
    try:
        logger.info("Initializing quality evaluator...")
        evaluator = QualityEvaluator()
        logger.info("Quality evaluator initialized successfully")
    except Exception as e:
        error_msg = f"Failed to initialize quality evaluator: {e}"
        logger.error(error_msg)
        errors.append(error_msg)
        evaluator = None
    
    if errors:
        st.session_state.initialization_errors = errors
        logger.error(f"Initialization completed with {len(errors)} errors")
    else:
        logger.info("All components initialized successfully")
    
    return vector_store_manager, style_manager, generator, evaluator


def main():
    """Main Streamlit application."""
    st.title("🌱 Erguvan AI Content Generator")
    st.subheader("Professional sustainability content with AI-powered generation")
    
    # Validate environment first
    env_issues = validate_environment()
    if env_issues:
        st.error("Environment validation failed:")
        for issue in env_issues:
            st.error(f"• {issue}")
    
    # Register cleanup handler
    if 'cleanup_registered' not in st.session_state:
        import atexit
        
        def cleanup_sessions():
            """Cleanup sessions on app shutdown."""
            if hasattr(st.session_state, 'session_manager') and st.session_state.session_manager:
                try:
                    asyncio.run(st.session_state.session_manager.shutdown())
                except Exception as e:
                    logger.error(f"Error during session cleanup: {e}")
        
        atexit.register(cleanup_sessions)
        st.session_state.cleanup_registered = True
    
    if env_issues:
        st.info("Please resolve these issues before continuing.")
        return
    
    # Initialize components
    if st.session_state.vector_store_manager is None:
        with st.spinner("Initializing system components..."):
            components = initialize_components()
            (st.session_state.vector_store_manager, 
             st.session_state.style_manager,
             st.session_state.generator,
             st.session_state.evaluator) = components
            
            # Show initialization status
            if st.session_state.initialization_errors:
                st.warning("System initialized with some issues:")
                for error in st.session_state.initialization_errors:
                    st.warning(f"• {error}")
                st.info("Some features may be limited.")
            else:
                st.success("System initialized successfully!")
            
            # Show which components are available
            available_components = []
            if st.session_state.vector_store_manager:
                available_components.append("Vector Store")
            if st.session_state.style_manager:
                available_components.append("Style Manager")
            if st.session_state.generator:
                available_components.append("Content Generator")
            if st.session_state.evaluator:
                available_components.append("Quality Evaluator")
            
            if available_components:
                st.info(f"Available components: {', '.join(available_components)}")
    
    # Sidebar navigation
    st.sidebar.header("Navigation")
    page = st.sidebar.selectbox(
        "Choose a page",
        ["Content Generation", "Document Management", "System Statistics", "Settings"]
    )
    
    if page == "Content Generation":
        content_generation_page()
    elif page == "Document Management":
        document_management_page()
    elif page == "System Statistics":
        statistics_page()
    elif page == "Settings":
        settings_page()


def content_generation_page():
    """Content generation interface with dynamic sample upload."""
    st.header("📝 Content Generation with Dynamic Samples")
    
    # Initialize session state
    if 'session_manager' not in st.session_state:
        try:
            from src.session.session_manager import SessionManager
            # Use the already initialized components
            if st.session_state.vector_store_manager and st.session_state.style_manager:
                st.session_state.session_manager = SessionManager(
                    st.session_state.vector_store_manager, 
                    st.session_state.style_manager
                )
            else:
                st.error("Session manager requires vector store and style manager to be initialized first.")
                return
        except Exception as e:
            st.error(f"Failed to initialize session manager: {e}")
            return
    
    if 'current_session' not in st.session_state:
        st.session_state.current_session = None
    
    # Step 1: Upload samples for this specific content generation
    st.subheader("1. Upload Topic-Specific Samples")
    st.info("💡 Upload the best examples related to your topic for higher quality content generation.")
    
    uploaded_files = st.file_uploader(
        "Upload sample documents (PDF, DOCX, PPTX, MD, TXT)",
        accept_multiple_files=True,
        type=['pdf', 'docx', 'pptx', 'md', 'txt'],
        help="Upload documents that contain excellent examples of content similar to what you want to generate"
    )
    
    # Process uploaded files
    if uploaded_files and st.session_state.current_session is None:
        # Create progress tracking for sample processing
        sample_progress_container = st.container()
        sample_status_container = st.container()
        
        with sample_progress_container:
            st.subheader("📁 Sample Processing Progress")
            sample_progress_bar = st.progress(0)
            
        with sample_status_container:
            sample_status_placeholder = st.empty()
            
        try:
            # Step 1: Create session
            with sample_status_placeholder.container():
                st.info("🆕 **Step 1/4:** Creating new session...")
            sample_progress_bar.progress(1/4)
            
            session = st.session_state.session_manager.create_session()
            st.session_state.current_session = session
            st.success("✅ Session created successfully")
            
            # Step 2: Upload files
            with sample_status_placeholder.container():
                st.info(f"📤 **Step 2/4:** Uploading {len(uploaded_files)} files...")
            sample_progress_bar.progress(2/4)
            
            for i, uploaded_file in enumerate(uploaded_files):
                file_content = uploaded_file.read()
                asyncio.run(
                    st.session_state.session_manager.add_sample_to_session(
                        session.session_id, 
                        file_content, 
                        uploaded_file.name
                    )
                )
                # Update progress during upload
                partial_progress = 2/4 + (i + 1) / len(uploaded_files) * 0.25
                sample_progress_bar.progress(partial_progress)
                
            st.success(f"✅ Successfully uploaded {len(uploaded_files)} files")
            
            # Step 3: Process samples
            with sample_status_placeholder.container():
                st.info("🔄 **Step 3/4:** Processing and analyzing samples...")
            sample_progress_bar.progress(3/4)
            
            processing_stats = asyncio.run(
                st.session_state.session_manager.process_session_samples(session.session_id)
            )
            
            # Step 4: Complete
            with sample_status_placeholder.container():
                st.info("✅ **Step 4/4:** Processing complete!")
            sample_progress_bar.progress(1.0)
            
            # Display detailed processing results
            st.success(f"✅ Processed {processing_stats['processed_samples']}/{processing_stats['total_samples']} samples")
            st.info(f"📊 Created {processing_stats['total_chunks']} content chunks in {processing_stats['processing_time']:.2f}s")
            
            if processing_stats['errors']:
                st.warning("⚠️ Some files had processing errors:")
                for error in processing_stats['errors']:
                    st.write(f"- {error}")
            else:
                st.success("🎉 All samples processed successfully!")
                
        except Exception as e:
            with sample_status_placeholder.container():
                st.error(f"❌ Sample processing failed: {e}")
                st.info("💡 Try uploading different file formats or check file sizes")
                
                # Show detailed error info in expander
                with st.expander("🔍 Detailed Error Information"):
                    import traceback
                    st.code(traceback.format_exc())
                    
                    # Show file info
                    if uploaded_files:
                        st.write("**Files attempted to upload:**")
                        for f in uploaded_files:
                            st.write(f"- {f.name} ({f.size} bytes, type: {f.type})")
                            
            st.session_state.current_session = None
    
    # Show current session info
    if st.session_state.current_session:
        session_info = st.session_state.session_manager.get_session_info(st.session_state.current_session.session_id)
        st.success(f"📁 Session active with {len(session_info['samples'])} samples")
        
        # Option to clear session and start over
        if st.button("🔄 Clear Session & Upload New Samples"):
            try:
                asyncio.run(st.session_state.session_manager.cleanup_session(st.session_state.current_session.session_id))
                st.session_state.current_session = None
                st.success("Session cleared successfully!")
                st.experimental_rerun()
            except Exception as e:
                st.error(f"Error clearing session: {e}")

        # --- DEBUG PANEL: Show session chunks ---
        with st.expander("🛠️ Debug: Show session chunks in vector store"):
            if st.button("Show session chunks", key="show_session_chunks_btn"):
                try:
                    chunks = st.session_state.session_manager.debug_list_session_chunks(
                        st.session_state.current_session.session_id
                    )
                    if chunks:
                        st.write(f"Found {len(chunks)} chunks in session vector store:")
                        for i, chunk in enumerate(chunks):
                            st.markdown(f"**Chunk {i+1}:**")
                            st.code(chunk['content'])
                            st.json(chunk['metadata'])
                    else:
                        st.warning("No chunks found in session vector store.")
                except Exception as e:
                    st.error(f"Error retrieving session chunks: {e}")
    
    # Show message if no samples uploaded yet
    elif not st.session_state.current_session:
        st.info("👆 Please upload sample documents above to begin content generation.")
        st.write("**Why upload samples?**")
        st.write("- Higher quality content generation")
        st.write("- Topic-specific examples and styles")
        st.write("- Better understanding of your content needs")
        st.write("- Fresh, current examples for each generation")
    
    # Step 2: Content generation form (always show)
    st.subheader("2. Generate Content")
    
    # Show different message based on sample upload status
    if st.session_state.current_session and st.session_state.current_session.is_processed:
        st.success("✅ Using your uploaded samples for enhanced content generation")
    elif st.session_state.current_session:
        st.warning("⚠️ Samples uploaded but not processed yet. Content will use default knowledge base.")
    else:
        st.info("ℹ️ No samples uploaded. Content will use default knowledge base.")
    
    # Always show the form
    # Content generation form
    with st.form("content_generation_form"):
        col1, col2 = st.columns(2)
        
        with col1:
            topic = st.text_input(
                "Topic",
                placeholder="e.g., EU Carbon Border Adjustment Mechanism",
                help="The main topic for content generation"
            )
            
            audience = st.selectbox(
                "Target Audience",
                ["C-Suite & Boards", "Sustainability Leads / ESG Managers", "Bank & Fintech Partners"],
                help="Who is the primary audience for this content?"
            )
            
            content_type = st.selectbox(
                "Content Type",
                ["document", "blog_post"],
                help="Type of content to generate"
            )
            
            tone_style = st.selectbox(
                "Tone Style",
                ["expert_report", "pragmatic_advisory", "inspirational"],
                help="Tone style for the content",
                format_func=lambda x: {
                    "expert_report": "Expert Report (white-papers, policy briefs)",
                    "pragmatic_advisory": "Pragmatic Advisory (implementation guides, playbooks)",
                    "inspirational": "Inspirational (keynotes, thought-leadership)"
                }[x]
            )
            
            # Dynamic length recommendations based on audience and content type
            if content_type == "blog_post":
                length_min, length_max, default_length = 750, 1100, 900
            elif audience == "C-Suite & Boards":
                length_min, length_max, default_length = 100, 800, 600
            elif audience == "Sustainability Leads / ESG Managers":
                length_min, length_max, default_length = 1500, 3000, 2000
            elif audience == "Bank & Fintech Partners":
                length_min, length_max, default_length = 100, 1200, 800
            else:
                length_min, length_max, default_length = 100, 5000, 1200
            
            desired_length = st.slider(
                f"Desired Length (words) - Recommended: {length_min}-{length_max}",
                min_value=100,
                max_value=5000,
                value=default_length,
                step=100
            )
        
        with col2:
            language = st.selectbox(
                "Language",
                ["EN", "TR"],
                format_func=lambda x: "English" if x == "EN" else "Turkish"
            )
            
            style_override = st.text_input(
                "Style Override (optional)",
                placeholder="e.g., executive summary format",
                help="Optional style instructions"
            )
            
        generate_button = st.form_submit_button("🚀 Generate Content", type="primary")
        
    # Generate content
    if generate_button and topic and audience:
        # Check if generator is available
        if not st.session_state.generator:
            st.error("Content generator is not available. Please check system initialization.")
            return
        
        # Create progress tracking containers
        progress_container = st.container()
        status_container = st.container()
        
        with progress_container:
            st.subheader("🚀 Content Generation Progress")
            progress_bar = st.progress(0)
            
        with status_container:
            status_placeholder = st.empty()
            
        try:
            # Step 1: Initialize request
            with status_placeholder.container():
                st.info("🔧 **Step 1/8:** Creating content request...")
            progress_bar.progress(1/8)
            
            content_request = ContentRequest(
                topic=topic,
                audience=audience,
                desired_length=desired_length,
                language=language,
                style_override=style_override,
                content_type=content_type,
                tone_style=tone_style
                )
                    
            # Step 2: Get session vector store
            with status_placeholder.container():
                st.info("🗂️ **Step 2/8:** Checking for uploaded samples...")
            progress_bar.progress(2/8)
                    
            session_vector_store = None
            if st.session_state.current_session:
                session_vector_store = run_async_in_thread(
                    st.session_state.session_manager.get_session_vector_store(
                        st.session_state.current_session.session_id
                    )
                )
                st.success("✅ Using your uploaded samples for enhanced context")
            else:
                st.info("ℹ️ Using default knowledge base (no samples uploaded)")
                    
            # Step 3: Context retrieval
            with status_placeholder.container():
                st.info("🔍 **Step 3/8:** Retrieving relevant context...")
            progress_bar.progress(3/8)
                    
            if session_vector_store:
                context_chunks = run_async_in_thread(
                    session_vector_store.search_for_generation(
                        topic=topic,
                        audience=audience,
                        language=language,
                        max_chunks=5
                    )
                )
                st.success(f"✅ Found {len(context_chunks)} relevant context chunks from your samples")
            else:
                context_chunks = run_async_in_thread(
                    st.session_state.vector_store_manager.search_for_generation(
                        topic=topic,
                        audience=audience,
                        language=language,
                        max_chunks=5
                    )
                )
                if context_chunks:
                    st.success(f"✅ Found {len(context_chunks)} relevant context chunks from knowledge base")
                else:
                    st.warning("⚠️ No relevant context found - generating with general knowledge")

            # --- DEBUG: Show context chunk content and scores ---
            with st.expander("🛠️ Debug: Show retrieved context chunks and scores"):
                if context_chunks:
                    for i, chunk in enumerate(context_chunks):
                        st.markdown(f"**Chunk {i+1}:** (Relevance score: {chunk.relevance_score:.3f})")
                        st.code(chunk.content)
                        st.json(chunk.metadata)
                else:
                    st.info("No context chunks retrieved.")
                    
            # Step 4: Content generation
            with status_placeholder.container():
                st.info("🤖 **Step 4/8:** Generating content with AI...")
            progress_bar.progress(4/8)
                    
            generated_content = run_async_in_thread(
                st.session_state.generator.generate_content(
                    content_request, 
                    session_vector_store=session_vector_store
                )
            )
                    
            with status_placeholder.container():
                st.success(f"✅ Content generated! ({generated_content.word_count} words, {generated_content.generation_time_seconds:.2f}s)")
                st.info(f"📊 Model used: {generated_content.model_used}")
                if generated_content.context_chunks_used > 0:
                    st.info(f"📚 Context chunks used: {generated_content.context_chunks_used}")
                if generated_content.style_patterns:
                    st.info(f"🎨 Style patterns applied: {len(generated_content.style_patterns)}")
                    
            # Step 5: Content evaluation
            with status_placeholder.container():
                st.info("📊 **Step 5/8:** Evaluating content quality...")
            progress_bar.progress(5/8)
                    
            evaluation = None
            try:
                evaluation = run_async_in_thread(
                    st.session_state.evaluator.evaluate_content(generated_content, context_chunks)
                )
                st.success("✅ Content evaluation completed")
                        
            except Exception as eval_error:
                st.warning(f"⚠️ Content evaluation failed: {eval_error}")
                st.info("🔄 Falling back to basic content display")
                    
            # Step 6: Display results
            with status_placeholder.container():
                st.info("📝 **Step 6/8:** Formatting results...")
            progress_bar.progress(6/8)
                    
            if evaluation:
                display_generation_results(generated_content, evaluation)
            else:
                display_content_only(generated_content)
                    
            # Step 7: Prepare downloads
            with status_placeholder.container():
                st.info("📄 **Step 7/8:** Preparing download options...")
            progress_bar.progress(7/8)
                    
            display_download_options(generated_content, evaluation)
                    
            # Step 8: Complete
            with status_placeholder.container():
                st.info("✅ **Step 8/8:** Generation complete!")
            progress_bar.progress(1.0)
                    
            st.success("🎉 Content generation completed successfully!")
                    
        except Exception as e:
            with status_placeholder.container():
                st.error(f"❌ Content generation failed: {e}")
                st.info("💡 Try adjusting your parameters or uploading more relevant samples")


def display_content_only(generated_content):
    """Display generated content without evaluation results."""
    # Content display
    st.subheader("Generated Content")
    
    # Title and metadata
    st.markdown(f"### {generated_content.title}")
    st.markdown(f"**Author:** {generated_content.author}")
    st.markdown(f"**Word Count:** {generated_content.word_count:,}")
    st.markdown(f"**Language:** {generated_content.language.upper()}")
    
    # Sections
    for section in generated_content.sections:
        if section.heading:
            st.markdown(f"#### {section.heading}")
        st.markdown(section.body)
        st.markdown("---")
    
    # Key takeaways
    if generated_content.key_takeaways:
        st.subheader("Key Takeaways")
        for takeaway in generated_content.key_takeaways:
            st.markdown(f"• {takeaway}")


def display_download_options(generated_content, evaluation):
    """Display download options for generated content."""
    st.subheader("Download Options")
    
    col1, col2 = st.columns(2)
    
    with col1:
        try:
            # Generate DOCX content in memory
            import io
            from docx import Document
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            
            # Create Word document in memory
            doc = Document()
            
            # Add title
            title_paragraph = doc.add_heading(generated_content.title, 0)
            title_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Add metadata
            doc.add_paragraph(f"Author: {generated_content.author}")
            doc.add_paragraph(f"Created: {generated_content.created_utc.strftime('%Y-%m-%d %H:%M:%S')} UTC")
            doc.add_paragraph(f"Audience: {generated_content.audience}")
            doc.add_paragraph(f"Word Count: {generated_content.word_count:,}")
            doc.add_paragraph(f"Language: {generated_content.language.upper()}")
            
            # Add separator
            doc.add_page_break()
            
            # Add sections
            for section in generated_content.sections:
                if section.heading:
                    doc.add_heading(section.heading, 1)
                
                # Add body paragraphs
                for paragraph in section.body.split('\n\n'):
                    if paragraph.strip():
                        doc.add_paragraph(paragraph.strip())
            
            # Add key takeaways
            if generated_content.key_takeaways:
                doc.add_heading("Key Takeaways", 1)
                for takeaway in generated_content.key_takeaways:
                    doc.add_paragraph(takeaway, style='List Bullet')
            
            # Add citations
            if generated_content.citations:
                doc.add_heading("References", 1)
                for i, citation in enumerate(generated_content.citations, 1):
                    citation_text = f"{i}. {citation.source}"
                    if citation.url:
                        citation_text += f" ({citation.url})"
                    doc.add_paragraph(citation_text)
            
            # Save to bytes
            docx_buffer = io.BytesIO()
            doc.save(docx_buffer)
            docx_buffer.seek(0)
            
            # Generate filename
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            filename = f"erguvan_content_{timestamp}.docx"
            
            # Download button
            st.download_button(
                label="📄 Download as DOCX",
                data=docx_buffer.getvalue(),
                file_name=filename,
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )
            
        except Exception as e:
            st.error(f"Failed to prepare DOCX: {e}")
    
    with col2:
        if evaluation is not None:
            try:
                # Generate evaluation report JSON
                evaluation_data = json.dumps(evaluation.to_dict(), indent=2)
                
                # Generate filename
                timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
                filename = f"erguvan_evaluation_{timestamp}.json"
                
                # Download button
                st.download_button(
                    label="📊 Download Evaluation Report",
                    data=evaluation_data,
                    file_name=filename,
                    mime="application/json"
                )
                
            except Exception as e:
                st.error(f"Failed to prepare evaluation report: {e}")
        else:
            st.info("Evaluation report not available (evaluation failed)")


def display_generation_results(generated_content, evaluation):
    """Display content generation results."""
    st.success("Content generated successfully!")
    
    # Quality evaluation summary
    col1, col2, col3, col4 = st.columns(4)
    
    with col1:
        st.metric(
            "Overall Recommendation",
            evaluation.overall_recommendation.upper(),
            delta=None
        )
    
    with col2:
        st.metric(
            "Plagiarism Score",
            f"{evaluation.plagiarism_report.overall_score:.1f}%",
            delta=f"{'✅' if evaluation.plagiarism_report.overall_score < 10 else '⚠️'}"
        )
    
    with col3:
        st.metric(
            "Quality Score",
            f"{evaluation.quality_report.overall_score:.1f}%",
            delta=f"{'✅' if evaluation.quality_report.overall_score >= 70 else '⚠️'}"
        )
    
    with col4:
        st.metric(
            "Brand Voice Score",
            f"{evaluation.brand_voice_report.overall_score:.1f}%",
            delta=f"{'✅' if evaluation.brand_voice_report.overall_score >= 90 else '⚠️'}"
        )
    
    # Content display
    st.subheader("Generated Content")
    
    # Title and metadata
    st.markdown(f"### {generated_content.title}")
    st.markdown(f"**Author:** {generated_content.author}")
    st.markdown(f"**Word Count:** {generated_content.word_count:,}")
    st.markdown(f"**Language:** {generated_content.language.upper()}")
    
    # Sections
    for section in generated_content.sections:
        if section.heading:
            st.markdown(f"#### {section.heading}")
        st.markdown(section.body)
        st.markdown("---")
    
    # Key takeaways
    if generated_content.key_takeaways:
        st.subheader("Key Takeaways")
        for takeaway in generated_content.key_takeaways:
            st.markdown(f"• {takeaway}")
    


def document_management_page():
    """Document management interface."""
    st.header("📚 Document Management")
    
    # File upload
    st.subheader("Upload Competitor Documents")
    
    uploaded_files = st.file_uploader(
        "Choose files",
        type=['pdf', 'docx', 'pptx', 'md'],
        accept_multiple_files=True,
        help="Upload competitor documents for analysis"
    )
    
    if uploaded_files:
        if st.button("🔄 Process Documents"):
            # Create progress containers
            progress_bar = st.progress(0)
            status_text = st.empty()
            
            try:
                # Use MarkItDown if available, otherwise fall back to DocumentLoader
                if MARKITDOWN_AVAILABLE:
                    loader = MarkItDownLoader()
                    loader_type = "MarkItDown (faster)"
                else:
                    loader = DocumentLoader()
                    loader_type = "DocumentLoader (standard)"
                    
                st.info(f"Using {loader_type} for document processing")
                processed_docs = []
                total_files = len(uploaded_files)
                
                for i, uploaded_file in enumerate(uploaded_files):
                    try:
                        # Update progress
                        progress = (i + 1) / total_files
                        progress_bar.progress(progress)
                        status_text.text(f"Processing {uploaded_file.name} ({i+1}/{total_files})...")
                        
                        # Validate file
                        if uploaded_file.size > 10 * 1024 * 1024:  # 10MB limit
                            st.warning(f"File {uploaded_file.name} is too large (>10MB), skipping")
                            continue
                        
                        # Save file with better error handling
                        samples_dir = Path("samples")
                        samples_dir.mkdir(exist_ok=True)
                        file_path = samples_dir / uploaded_file.name
                        
                        # Check if file already exists
                        if file_path.exists():
                            st.info(f"File {uploaded_file.name} already exists, overwriting")
                        
                        logger.info(f"Processing file: {uploaded_file.name}")
                        
                        with open(file_path, "wb") as f:
                            f.write(uploaded_file.getvalue())
                        
                        # Process document with MarkItDown (much faster!)
                        try:
                            def process_document():
                                content, metadata = loader.load_document(str(file_path))
                                chunks = loader.chunk_document(content, metadata)
                                return content, metadata, chunks
                            
                            # Run document processing in thread with shorter timeout (MarkItDown is much faster)
                            future = st.session_state.async_executor.submit(process_document)
                            
                            try:
                                # Adjust timeout based on loader type
                                timeout = 10 if MARKITDOWN_AVAILABLE else 30
                                content, metadata, chunks = future.result(timeout=timeout)
                                
                                processed_docs.append({
                                    'filename': uploaded_file.name,
                                    'word_count': metadata.word_count,
                                    'language': metadata.language,
                                    'chunks': len(chunks),
                                    'file_type': metadata.file_type
                                })
                                
                                logger.info(f"Successfully processed {uploaded_file.name}: {len(chunks)} chunks")
                                
                            except TimeoutError:
                                timeout = 10 if MARKITDOWN_AVAILABLE else 30
                                error_msg = f"Processing {uploaded_file.name} timed out (>{timeout}s)"
                                logger.error(error_msg)
                                st.warning(error_msg)
                                
                                # Clean up file if processing timed out
                                if file_path.exists():
                                    file_path.unlink()
                                
                        except Exception as process_error:
                            error_msg = f"Failed to process {uploaded_file.name}: {process_error}"
                            logger.error(error_msg)
                            st.warning(error_msg)
                            
                            # Clean up file if processing failed
                            if file_path.exists():
                                file_path.unlink()
                            
                    except Exception as file_error:
                        error_msg = f"Failed to handle file {uploaded_file.name}: {file_error}"
                        logger.error(error_msg)
                        st.error(error_msg)
                
                # Clear progress indicators
                progress_bar.empty()
                status_text.empty()
                
                # Display results
                if processed_docs:
                    success_msg = f"Processed {len(processed_docs)} documents successfully"
                    if MARKITDOWN_AVAILABLE:
                        success_msg += " with MarkItDown!"
                    st.success(success_msg)
                    
                    df = pd.DataFrame(processed_docs)
                    st.dataframe(df)
                    
                    # Show processing summary
                    total_chunks = sum(doc['chunks'] for doc in processed_docs)
                    total_words = sum(doc['word_count'] for doc in processed_docs)
                    
                    col1, col2, col3 = st.columns(3)
                    with col1:
                        st.metric("Documents Processed", len(processed_docs))
                    with col2:
                        st.metric("Total Chunks", total_chunks)
                    with col3:
                        st.metric("Total Words", f"{total_words:,}")
                        
                    # Show file type distribution
                    file_types = [doc['file_type'] for doc in processed_docs]
                    type_counts = pd.Series(file_types).value_counts()
                    
                    st.subheader("File Type Distribution")
                    st.bar_chart(type_counts)
                    
                else:
                    st.warning("No documents were processed successfully")
                    
            except Exception as e:
                logger.error(f"Document processing failed: {e}")
                st.error(f"Document processing failed: {e}")
            
            finally:
                # Clean up progress indicators
                try:
                    progress_bar.empty()
                    status_text.empty()
                except:
                    pass
    
    # Document search
    st.subheader("Search Documents")
    
    search_query = st.text_input("Search query", placeholder="e.g., carbon pricing")
    
    if search_query:
        # Check if vector store is available
        if not st.session_state.vector_store_manager:
            st.error("Vector store is not available. Please check system initialization.")
            return
        
        try:
            results = run_async_in_thread(
                st.session_state.vector_store_manager.vector_store.search_chunks(
                    query=search_query,
                    limit=10
                )
            )
            
            if results:
                st.write(f"Found {len(results)} results:")
                
                for i, result in enumerate(results):
                    with st.expander(f"Result {i+1} - Relevance: {result.relevance_score:.3f}"):
                        st.write(f"**Document:** {result.metadata.get('file_name', 'Unknown')}")
                        if result.heading:
                            st.write(f"**Heading:** {result.heading}")
                        st.write(f"**Content:** {result.content}")
            else:
                st.info("No results found")
                
        except Exception as e:
            st.error(f"Search failed: {e}")


def statistics_page():
    """System statistics interface."""
    st.header("📊 System Statistics")
    
    if not st.session_state.vector_store_manager:
        st.error("Vector store is not available. Please check system initialization.")
        return
    
    try:
        stats = st.session_state.vector_store_manager.get_stats()
        
        # Overview metrics
        col1, col2, col3, col4 = st.columns(4)
        
        with col1:
            st.metric("Total Documents", f"{stats.total_documents:,}")
        
        with col2:
            st.metric("Total Chunks", f"{stats.total_chunks:,}")
        
        with col3:
            st.metric("Index Size", f"{stats.index_size_mb:.1f} MB")
        
        with col4:
            st.metric("Languages", len(stats.languages))
        
        # Detailed information
        st.subheader("Detailed Statistics")
        
        info_data = {
            "Metric": [
                "Total Embeddings",
                "Average Chunk Size",
                "Languages",
                "Last Updated"
            ],
            "Value": [
                f"{stats.total_embeddings:,}",
                f"{stats.average_chunk_size:.1f} characters",
                ", ".join(stats.languages) if stats.languages else "None",
                stats.last_updated.strftime("%Y-%m-%d %H:%M:%S")
            ]
        }
        
        df = pd.DataFrame(info_data)
        st.dataframe(df, use_container_width=True)
        
    except Exception as e:
        st.error(f"Failed to load statistics: {e}")


def settings_page():
    """Settings interface."""
    st.header("⚙️ Settings")
    
    # Configuration display
    col1, col2 = st.columns([3, 1])
    with col1:
        st.subheader("Current Configuration")
    with col2:
        if st.button("🔄 Reload Config"):
            config_manager.reload_config()
            st.success("Configuration reloaded!")
    
    try:
        model_config = config_manager.get_model_config()
        brand_config = config_manager.get_brand_config()
        
        col1, col2 = st.columns(2)
        
        with col1:
            st.markdown("**Model Configuration:**")
            st.write(f"Provider: {model_config.provider}")
            st.write(f"Generation Model: {model_config.generation_model}")
            st.write(f"Analysis Model: {model_config.analysis_model}")
            st.write(f"Embedding Model: {model_config.embedding_model}")
            st.write(f"Temperature: {model_config.temperature}")
            st.write(f"Max Tokens: {model_config.max_tokens}")
            st.write(f"Max Cost per Request: ${model_config.max_cost_per_request}")
        
        with col2:
            st.markdown("**Brand Configuration:**")
            st.write(f"Primary Tone: {brand_config.primary_tone}")
            st.write(f"Flesch-Kincaid Grade Max: {brand_config.flesch_score_min}")
            st.write(f"Sentence Length: {brand_config.preferred_sentence_length}")
            st.write(f"Secondary Traits: {', '.join(brand_config.secondary_traits[:3])}...")
            st.write(f"Avoid: {', '.join(brand_config.avoid[:2])}...")
        
    except Exception as e:
        st.error(f"Failed to load configuration: {e}")
    
    # Environment status
    st.subheader("Environment Status")
    
    api_key = os.getenv('OPENAI_API_KEY')
    if api_key:
        st.success("✅ OpenAI API Key configured")
    else:
        st.error("❌ OpenAI API Key not found")
    
    # System health
    st.subheader("System Health")
    
    if st.session_state.vector_store_manager:
        try:
            health = st.session_state.vector_store_manager.health_check()
            
            if health["healthy"]:
                st.success("✅ System healthy")
            else:
                st.error("❌ System issues detected")
                
            st.json(health)
            
        except Exception as e:
            st.error(f"Health check failed: {e}")
    else:
        st.error("Vector store is not available for health check.")


if __name__ == "__main__":
    main()